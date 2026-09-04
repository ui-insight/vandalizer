"""Tests for workflow engine pure functions — step name sanitization,
HTML text extraction, output formatting, and node contracts."""

import base64
import json

from app.services.workflow_engine import (
    DocumentNode,
    AddDocumentNode,
    DataExportNode,
    DocumentRendererNode,
    WorkflowEngine,
    _extract_text_from_html,
    _stringify_value,
    format_extraction_results,
    sanitize_step_name,
)


# ---------------------------------------------------------------------------
# sanitize_step_name — prevents MongoDB operator injection via . and $
# ---------------------------------------------------------------------------

class TestSanitizeStepName:
    def test_dots_replaced(self):
        assert sanitize_step_name("foo.bar") == "foo_bar"

    def test_dollars_replaced(self):
        assert sanitize_step_name("$where") == "where"

    def test_multiple_special_chars(self):
        assert sanitize_step_name("$foo.bar$") == "foo_bar"

    def test_whitespace_collapsed(self):
        assert sanitize_step_name("hello   world") == "hello_world"

    def test_double_underscores_collapsed(self):
        assert sanitize_step_name("step__name") == "step_name"

    def test_empty_string_returns_step(self):
        assert sanitize_step_name("") == "step"

    def test_only_special_chars_returns_step(self):
        assert sanitize_step_name("$.$") == "step"

    def test_whitespace_only_returns_step(self):
        assert sanitize_step_name("   ") == "step"

    def test_leading_trailing_underscores_stripped(self):
        assert sanitize_step_name("_name_") == "name"

    def test_normal_name_unchanged(self):
        assert sanitize_step_name("Extraction") == "Extraction"

    def test_tabs_and_newlines(self):
        assert sanitize_step_name("step\t\nname") == "step_name"


# ---------------------------------------------------------------------------
# _extract_text_from_html — strips dangerous tags, normalizes whitespace
# ---------------------------------------------------------------------------

class TestExtractTextFromHtml:
    def test_basic_text(self):
        assert _extract_text_from_html("<p>Hello world</p>") == "Hello world"

    def test_script_tags_removed(self):
        html = "<p>Text</p><script>alert('xss')</script><p>More</p>"
        result = _extract_text_from_html(html)
        assert "alert" not in result
        assert "Text" in result
        assert "More" in result

    def test_style_tags_removed(self):
        html = "<style>.foo{color:red}</style><p>Content</p>"
        result = _extract_text_from_html(html)
        assert "color" not in result
        assert "Content" in result

    def test_nav_footer_header_removed(self):
        html = "<nav>NavBar</nav><main>Main Content</main><footer>Footer</footer>"
        result = _extract_text_from_html(html)
        assert "NavBar" not in result
        assert "Footer" not in result
        assert "Main Content" in result

    def test_form_tags_removed(self):
        html = "<form><input type='text' value='secret'/></form><p>Visible</p>"
        result = _extract_text_from_html(html)
        assert "secret" not in result
        assert "Visible" in result

    def test_whitespace_normalized(self):
        html = "<p>  lots   of    spaces  </p>"
        result = _extract_text_from_html(html)
        assert "  " not in result  # no double spaces
        assert "lots of spaces" in result

    def test_excessive_newlines_collapsed(self):
        html = "<p>A</p><p></p><p></p><p></p><p>B</p>"
        result = _extract_text_from_html(html)
        assert "\n\n\n" not in result

    def test_empty_html(self):
        assert _extract_text_from_html("") == ""

    def test_aside_removed(self):
        html = "<aside>Sidebar</aside><article>Article</article>"
        result = _extract_text_from_html(html)
        assert "Sidebar" not in result
        assert "Article" in result


# ---------------------------------------------------------------------------
# format_extraction_results / _stringify_value
# ---------------------------------------------------------------------------

class TestStringifyValue:
    def test_none_returns_na(self):
        assert _stringify_value(None) == "N/A"

    def test_string_passthrough(self):
        assert _stringify_value("hello") == "hello"

    def test_int_to_string(self):
        assert _stringify_value(42) == "42"

    def test_list_joined(self):
        assert _stringify_value(["a", "b", "c"]) == "a, b, c"

    def test_list_with_none_filtered(self):
        assert _stringify_value(["a", None, "b"]) == "a, b"

    def test_dict_to_json(self):
        result = _stringify_value({"key": "val"})
        assert json.loads(result) == {"key": "val"}


class TestFormatExtractionResults:
    def test_none_returns_empty(self):
        assert format_extraction_results(None) == ""

    def test_single_dict(self):
        result = format_extraction_results({"Name": "Alice", "Age": "30"})
        assert "**Name**" in result
        assert "Alice" in result
        assert "**Age**" in result
        assert "30" in result

    def test_list_of_dicts(self):
        result = format_extraction_results([
            {"Name": "Alice"},
            {"Name": "Bob"},
        ])
        assert "Result 1" in result
        assert "Result 2" in result
        assert "Alice" in result
        assert "Bob" in result

    def test_single_item_list_no_result_header(self):
        result = format_extraction_results([{"Name": "Alice"}])
        assert "Result" not in result
        assert "Alice" in result

    def test_scalar_value(self):
        result = format_extraction_results("just a string")
        assert result == "just a string"

    def test_empty_list(self):
        result = format_extraction_results([])
        assert result == ""


# ---------------------------------------------------------------------------
# WorkflowEngine._format_final_output
# ---------------------------------------------------------------------------

class TestFormatFinalOutput:
    def setup_method(self):
        self.engine = WorkflowEngine()

    def test_none_returns_empty_string(self):
        assert self.engine._format_final_output(None) == ""

    def test_string_passthrough(self):
        assert self.engine._format_final_output("hello") == "hello"

    def test_int_converted_to_string(self):
        assert self.engine._format_final_output(42) == "42"

    def test_single_item_dict_list_unwrapped(self):
        result = self.engine._format_final_output([{"key": "val"}])
        parsed = json.loads(result)
        assert parsed == {"key": "val"}

    def test_multi_item_list_gets_headers(self):
        result = self.engine._format_final_output(["first", "second"])
        assert "### Result 1" in result
        assert "### Result 2" in result

    def test_file_download_dict_passthrough(self):
        download = {"type": "file_download", "data_b64": "abc", "filename": "out.csv"}
        result = self.engine._format_final_output(download)
        assert result == download

    def test_regular_dict_to_json(self):
        result = self.engine._format_final_output({"key": "val"})
        assert json.loads(result) == {"key": "val"}

    def test_empty_list(self):
        assert self.engine._format_final_output([]) == ""

    def test_list_item_starting_with_heading_no_extra_header(self):
        result = self.engine._format_final_output(["# Title", "other"])
        assert "### Result" not in result.split("# Title")[0]  # Title not re-wrapped


# ---------------------------------------------------------------------------
# Node contracts — DocumentNode, AddDocumentNode, DataExportNode
# ---------------------------------------------------------------------------

class TestDocumentNode:
    def test_output_contains_uuids(self):
        node = DocumentNode({"doc_uuids": ["uuid1", "uuid2"]})
        result = node.process()
        assert result["output"] == ["uuid1", "uuid2"]
        assert result["step_name"] == "Document"
        assert result["input"] is None

    def test_empty_uuids(self):
        node = DocumentNode({})
        result = node.process()
        assert result["output"] == []


class TestAddDocumentNode:
    def test_joins_doc_texts(self):
        node = AddDocumentNode({"doc_texts": ["Hello", "World"]})
        result = node.process({"output": None})
        assert result["output"] == "Hello\nWorld"
        assert result["step_name"] == "AddDocument"

    def test_empty_texts_is_a_step_error(self):
        """The same guard Add Website carries: a document-attachment step with
        nothing to attach used to return "" and let the run finish Completed."""
        node = AddDocumentNode({})
        result = node.process({"output": None})
        assert result["output"] == ""
        assert "no document text" in result["error"]

    def test_whitespace_only_texts_are_a_step_error(self):
        node = AddDocumentNode({"doc_texts": ["  ", "\n"]})
        result = node.process({"output": None})
        assert "no document text" in result["error"]


class TestDataExportNode:
    def test_json_export(self):
        node = DataExportNode({"format": "json", "filename": "test"})
        result = node.process({"output": {"key": "value"}})
        output = result["output"]
        assert output["type"] == "file_download"
        assert output["file_type"] == "json"
        assert output["filename"] == "test.json"
        decoded = base64.b64decode(output["data_b64"]).decode()
        assert json.loads(decoded) == {"key": "value"}

    def test_csv_export_list_of_dicts(self):
        data = [{"name": "Alice", "age": "30"}, {"name": "Bob", "age": "25"}]
        node = DataExportNode({"format": "csv", "filename": "people"})
        result = node.process({"output": data})
        output = result["output"]
        assert output["file_type"] == "csv"
        decoded = base64.b64decode(output["data_b64"]).decode()
        assert "name" in decoded  # header row
        assert "Alice" in decoded
        assert "Bob" in decoded

    def test_csv_export_single_dict(self):
        node = DataExportNode({"format": "csv", "filename": "single"})
        result = node.process({"output": {"a": "1", "b": "2"}})
        decoded = base64.b64decode(result["output"]["data_b64"]).decode()
        assert "a" in decoded
        assert "1" in decoded


class TestDocumentRendererNode:
    def test_markdown_render(self):
        node = DocumentRendererNode({"format": "md", "filename": "report"})
        result = node.process({"output": "# Hello"})
        output = result["output"]
        assert output["type"] == "file_download"
        assert output["file_type"] == "md"
        assert output["filename"] == "report.md"
        decoded = base64.b64decode(output["data_b64"]).decode()
        assert decoded == "# Hello"

    def test_dict_input_serialized_to_json(self):
        node = DocumentRendererNode({"format": "txt", "filename": "out"})
        result = node.process({"output": {"k": "v"}})
        decoded = base64.b64decode(result["output"]["data_b64"]).decode()
        assert json.loads(decoded) == {"k": "v"}

    # The step advertised "DOCX, PDF, etc." and produced .md or .txt whatever
    # the setting said (support ticket). It now renders the formats it names.

    def test_pdf_render_produces_a_pdf_from_markdown(self):
        node = DocumentRendererNode({"format": "pdf", "filename": "award_summary"})
        result = node.process({"output": "# Award\n\n**PI:** Dr. Ada\n\n- one\n- two"})
        out = result["output"]
        assert out["file_type"] == "pdf" and out["filename"] == "award_summary.pdf"
        pdf = base64.b64decode(out["data_b64"])
        assert pdf.startswith(b"%PDF")
        import fitz

        text = "".join(page.get_text() for page in fitz.open(stream=pdf, filetype="pdf"))
        assert "award summary" in text  # filename doubles as the document title
        assert "Dr. Ada" in text and "one" in text
        assert "**" not in text and "# Award" not in text  # markdown rendered, not pasted

    def test_pdf_title_can_be_set(self):
        node = DocumentRendererNode({"format": "pdf", "filename": "x", "title": "Quarterly Report"})
        pdf = base64.b64decode(node.process({"output": "body"})["output"]["data_b64"])
        import fitz

        assert "Quarterly Report" in fitz.open(stream=pdf, filetype="pdf")[0].get_text()

    def test_docx_render_produces_a_word_document(self):
        node = DocumentRendererNode({"format": "docx", "filename": "letter"})
        result = node.process({"output": "# Dear PI\n\nYour award is **approved**."})
        out = result["output"]
        assert out["file_type"] == "docx" and out["filename"] == "letter.docx"
        raw = base64.b64decode(out["data_b64"])
        assert raw[:2] == b"PK"
        import io

        from docx import Document

        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs]
        assert "Dear PI" in paragraphs
        assert any("approved" in p and "**" not in p for p in paragraphs)

    def test_structured_input_renders_as_a_table_in_pdf_and_docx(self):
        rows = [{"field": "rate", "value": "47.5%"}, {"field": "cap", "value": "n/a"}]
        pdf = base64.b64decode(DocumentRendererNode({"format": "pdf"}).process({"output": rows})["output"]["data_b64"])
        import fitz

        assert "47.5%" in fitz.open(stream=pdf, filetype="pdf")[0].get_text()
        import io

        from docx import Document

        raw = base64.b64decode(DocumentRendererNode({"format": "docx"}).process({"output": rows})["output"]["data_b64"])
        assert Document(io.BytesIO(raw)).tables[0].rows[1].cells[1].text == "47.5%"

    def test_unknown_format_falls_back_to_text(self):
        result = DocumentRendererNode({"format": "rtf", "filename": "f"}).process({"output": "x"})
        assert result["output"]["filename"] == "f.txt"

    def test_blank_filename_defaults(self):
        result = DocumentRendererNode({"format": "md", "filename": "  "}).process({"output": "x"})
        assert result["output"]["filename"] == "output.md"

    def test_file_input_is_a_configuration_error(self):
        upstream = {"type": "file_download", "data_b64": "QUJD", "file_type": "pdf", "filename": "form-filled.pdf"}
        result = DocumentRendererNode({"format": "pdf"}).process({"output": upstream})
        assert "received a file (form-filled.pdf)" in result["error"]
        assert result["output"] == ""


class TestKnowledgeBaseQueryNodeApproximatePages:
    """A KB Query node in ``answer`` mode writes prose that flows into
    downstream steps and exports, where the ``page_approximate`` flag on
    ``sources`` does not reach. If the model restates an estimated page as
    exact, that is the last word anyone sees.
    """

    def _run_answer_mode(self, monkeypatch, *, approximate: bool):
        from unittest.mock import MagicMock

        from app.services import workflow_engine as we

        captured = {}

        def fake_llm_chat_model(*, prompt, data, **kwargs):
            captured["prompt"] = prompt
            captured["data"] = data
            return "answer text"

        chunk = {
            "content": "Indirect costs are capped at 58% MTDC.",
            "metadata": {
                "source_name": "PAPPG.pdf",
                "source_id": "src-1",
                "page": 234,
                "page_approximate": approximate,
            },
            "chunk_id": "c1",
            "score": 0.1,
            "similarity": 0.9,
        }

        dm = MagicMock()
        dm.query_kb.return_value = [chunk]
        # The node imports DocumentManager inside the function, so it has to be
        # patched where it is defined rather than on workflow_engine.
        from app.services import document_manager as dm_mod
        monkeypatch.setattr(dm_mod, "DocumentManager", lambda *a, **k: dm)
        monkeypatch.setattr(we, "llm_chat_model", fake_llm_chat_model)

        node = we.KnowledgeBaseQueryNode({
            "kb_uuid": "kb-1", "query": "indirect rate?", "mode": "answer",
        })
        result = node.process({"output": None})
        return captured, result

    def test_the_tilde_is_explained_when_a_page_is_estimated(self, monkeypatch):
        captured, result = self._run_answer_mode(monkeypatch, approximate=True)

        assert "p. ~234" in captured["data"], (
            "the passage label lost its hedge before reaching the model"
        )
        assert "estimate" in captured["prompt"], (
            "the model was shown a tilde with nothing explaining it, and the "
            "instruction's own example is an un-hedged page"
        )
        assert "explicitly" in captured["prompt"], (
            "restating an approximate page as exact was not ruled out by name"
        )
        assert result["retrieved_sources"][0]["page_approximate"] is True

    def test_no_estimated_page_means_no_rule_about_them(self, monkeypatch):
        captured, _ = self._run_answer_mode(monkeypatch, approximate=False)

        assert "p. 234" in captured["data"]
        assert "p. ~234" not in captured["data"]
        assert "estimate" not in captured["prompt"]

    def test_the_citation_example_does_not_teach_dropping_the_tilde(self):
        """The instruction shows `[PAPPG.pdf · p. 234]` as the citation shape.
        A model normalising to that example is how a hedged label silently
        becomes an exact one, so the instruction must say to copy the locator
        rather than to match the example's form.
        """
        from app.services.workflow_engine import KB_ANSWER_INSTRUCTION

        assert "copying the locator exactly as shown" in KB_ANSWER_INSTRUCTION
