"""Markdown/data → Word (.docx) rendering.

Shared by the workflow download endpoint and the save-to-folder / passive
storage paths so every export surface produces the same branded Word
documents. python-docx itself is imported lazily inside the render functions.
"""

import io
import json
import re

_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_INLINE_CODE = re.compile(r"`(.+?)`")
_MD_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _add_runs_with_formatting(paragraph, text: str) -> None:
    """Add runs to a docx paragraph, parsing inline **bold**, *italic*, `code`, [link](url)."""
    # Drop markdown link syntax — keep the visible text only (URLs in LLM output
    # are often hallucinated and useless in a Word doc).
    text = _MD_LINK.sub(r"\1", text)

    pos = 0
    # Combined scan: find the next inline marker and emit the preceding plain run.
    while pos < len(text):
        next_match = None
        next_kind = None
        for kind, pattern in (("bold", _INLINE_BOLD), ("code", _INLINE_CODE), ("italic", _INLINE_ITALIC)):
            m = pattern.search(text, pos)
            if m and (next_match is None or m.start() < next_match.start()):
                next_match = m
                next_kind = kind
        if next_match is None:
            paragraph.add_run(text[pos:])
            return
        if next_match.start() > pos:
            paragraph.add_run(text[pos:next_match.start()])
        run = paragraph.add_run(next_match.group(1))
        if next_kind == "bold":
            run.bold = True
        elif next_kind == "italic":
            run.italic = True
        elif next_kind == "code":
            run.font.name = "Consolas"
        pos = next_match.end()


# Separator row of a GitHub-flavored markdown table (e.g. ``|---|:--:|-|``).
# One-or-more dashes per column matches the GFM spec (and remark-gfm, which the
# UI renders with), so a model emitting ``|--|--|`` still yields a real table.
# Kept identical to pdf_service so DOCX and PDF detect the same tables.
_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
# Horizontal rule: 3+ repeats of the same -, * or _ (e.g. ``---``). Mirrors
# pdf_service so a standalone rule renders as a line, not literal text.
_MD_HR_RE = re.compile(r"^\s{0,3}([-*_])(\s*\1){2,}\s*$")


def _split_md_table_row(line: str) -> list[str]:
    """Split a markdown table row into trimmed cell strings."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # Honor backslash-escaped pipes by stashing them behind a placeholder.
    line = line.replace(r"\|", "\x00")
    return [c.strip().replace("\x00", "|") for c in line.split("|")]


# Downloaded-table theme — matches the gold PDF header in pdf_service so Word and
# PDF downloads share the Vandalizer brand look (UI highlight color #eab308).
_DOCX_TABLE_HEADER_FILL = "EAB308"   # brand gold header fill
_DOCX_TABLE_STRIPE_FILL = "FDF9EB"   # faint gold tint for zebra striping
_DOCX_TABLE_BORDER = "E5E7EB"        # light-gray grid
_DOCX_HEADING_HEX = "191919"         # charcoal headings (override Word's blue)


def _shade_docx_cell(cell, fill_hex: str) -> None:
    """Set a table cell's background fill (python-docx has no high-level API)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_docx_table_borders(table, color_hex: str) -> None:
    """Apply a thin single-line grid border (color_hex) to every table edge."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")       # 4 eighths-of-a-point = 0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _add_docx_hr(doc) -> None:
    """Render a markdown horizontal rule as an empty paragraph with a thin,
    light-gray bottom border (Word has no native ``<hr>``)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _DOCX_TABLE_BORDER)
    borders.append(bottom)
    p_pr.append(borders)


def _apply_docx_table_theme(table) -> None:
    """Apply the Vandalizer gold table theme to a Word table.

    Gold (#eab308) header row with bold black text, faint gold zebra striping on
    body rows, and a light-gray grid. Replaces the built-in (blue) ``Light Grid
    Accent 1`` style so DOCX downloads match the gold-headed PDF tables.
    """
    table.style = "Table Grid"
    _set_docx_table_borders(table, _DOCX_TABLE_BORDER)
    rows = table.rows
    if not rows:
        return
    for cell in rows[0].cells:
        _shade_docx_cell(cell, _DOCX_TABLE_HEADER_FILL)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for idx, row in enumerate(rows[1:], start=1):
        if idx % 2 == 0:  # zebra-stripe every other body row
            for cell in row.cells:
                _shade_docx_cell(cell, _DOCX_TABLE_STRIPE_FILL)


def _add_markdown_table_to_docx(doc, headers: list[str], rows: list[list[str]]) -> None:
    """Render parsed markdown-table cells as a real Word table.

    Uses the shared gold table theme (see ``_apply_docx_table_theme``), with
    inline **bold**/*italic*/`code` honored inside every cell.
    """
    col_count = max(len(headers), max((len(r) for r in rows), default=0)) or 1
    headers = headers + [""] * (col_count - len(headers))

    table = doc.add_table(rows=1, cols=col_count)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _add_runs_with_formatting(hdr_cells[i].paragraphs[0], h)

    for row in rows:
        row = row + [""] * (col_count - len(row))
        cells = table.add_row().cells
        for i in range(col_count):
            _add_runs_with_formatting(cells[i].paragraphs[0], row[i])

    _apply_docx_table_theme(table)


def markdown_to_docx(text: str):
    """Render a markdown-ish string into a python-docx Document.

    Handles ATX headings, unordered/ordered lists, GitHub-flavored tables,
    blank-line paragraphs, and inline bold/italic/code. Unknown syntax falls
    back to a plain paragraph.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    heading_color = RGBColor.from_string(_DOCX_HEADING_HEX)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = (text or "").splitlines()
    n = len(lines)
    i = 0
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Markdown table — a header row followed by a dash-separator row. Consume
        # the header, separator, and all contiguous body rows into a real table.
        if "|" in line and i + 1 < n and _MD_TABLE_SEP_RE.match(lines[i + 1].rstrip()):
            headers = _split_md_table_row(line)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n:
                body = lines[i].rstrip()
                if not body.strip() or "|" not in body:
                    break
                rows.append(_split_md_table_row(body))
                i += 1
            _add_markdown_table_to_docx(doc, headers, rows)
            continue

        # Horizontal rule (``---`` / ``***`` / ``___``) → a thin border, not text.
        if _MD_HR_RE.match(line):
            _add_docx_hr(doc)
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            h = doc.add_heading(heading.group(2).strip(), level=level)
            for run in h.runs:
                run.font.color.rgb = heading_color  # charcoal, not Word's blue
            i += 1
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_formatting(p, bullet.group(1).strip())
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_formatting(p, numbered.group(1).strip())
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs_with_formatting(p, line)
        i += 1

    return doc


def data_to_docx_bytes(data) -> bytes:
    """Serialize workflow output (str/dict/list) to a .docx byte payload."""
    from docx import Document
    from docx.shared import Inches, Pt

    if isinstance(data, str):
        try:
            parsed = json.loads(data)
        except (json.JSONDecodeError, ValueError):
            parsed = None
        if isinstance(parsed, (dict, list)):
            data = parsed

    if isinstance(data, str):
        doc = markdown_to_docx(data)
    elif isinstance(data, list) and data and isinstance(data[0], dict):
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Inches(1)
            section.left_margin = section.right_margin = Inches(1)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)
        headers = list(dict.fromkeys(k for row in data for k in row.keys()))
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
        for row in data:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                val = row.get(h, "")
                cells[i].text = json.dumps(val, default=str) if isinstance(val, (dict, list)) else str(val if val is not None else "")
        _apply_docx_table_theme(table)
    elif isinstance(data, dict):
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Inches(1)
            section.left_margin = section.right_margin = Inches(1)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"
        for k, v in data.items():
            cells = table.add_row().cells
            cells[0].text = str(k)
            cells[1].text = json.dumps(v, default=str) if isinstance(v, (dict, list)) else str(v)
        _apply_docx_table_theme(table)
    elif isinstance(data, list):
        doc = markdown_to_docx("\n".join(f"- {item}" for item in data))
    else:
        doc = markdown_to_docx("" if data is None else str(data))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
