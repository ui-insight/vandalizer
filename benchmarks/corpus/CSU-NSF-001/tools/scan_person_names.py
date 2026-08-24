"""Independent identity check: find natural-person names outside permitted places.

The corpus ships `validate_identity_safety.py`, which is a *denylist* of the
specific names v0.3.2 used. That proves the old names are gone; it cannot catch
a new real-person name, which is the failure mode that started this. This is
the structural check instead: find everything that *looks* like a person's
name anywhere in what it is pointed at, then subtract the places where a real
name is legitimate (References Cited, figure credits). Whatever is left is a
finding.

Over-reporting is the intended bias — every hit gets read by a human once.
Which is also why a bare exit code cannot gate this in CI: the clean corpus
carries a stable set of reviewed candidates (invented place names, mostly), so
"any finding fails" fails always. `--baseline` is the reviewed set: the run
fails on findings that are *not* in it, and says so about each one. Stale
baseline entries are reported but do not fail — a name disappearing is the
direction we want.

Run: cd backend && uv run python \\
       ../benchmarks/corpus/CSU-NSF-001/tools/scan_person_names.py DIR \\
       [--baseline ../benchmarks/corpus/CSU-NSF-001/tools/name_scan_baseline_tree.json]

To regenerate a baseline, `--emit-baseline PATH` writes the current findings in
baseline form. Read every line it writes before committing it: an unreviewed
baseline is an unreviewed corpus, and the whole point of the file is that a new
name shows up as an added line in the diff.
"""
import argparse
import fnmatch
import json
import re
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

# Documents where a real person's name is legitimate: genuine scholarly
# attribution. Everything else in the package must be role identifiers.
PERMITTED_FILES = {"07_References_Cited.pdf", "07_References_Cited.docx"}

# In-tree files that necessarily contain the retired real names, matched by
# relative-path suffix. The denylist has to spell them out to test for them,
# and a baseline has to record what was reviewed; both exemptions are visible
# in the same diff as any corpus change.
PERMITTED_TREE_FILES = {"tools/validate_identity_safety.py"}
PERMITTED_TREE_GLOBS = ("name_scan_baseline_*.json",)

# Lines carrying an open-license figure credit may name the creator.
CREDIT_MARKERS = ("CC BY", "CC-BY", "CC0", "Public domain", "photo credit",
                  "Photo credit", "Image credit", "image credit", "Adapted from",
                  "Source:", "Credit:", "credit:")

# Text files that are part of the corpus rather than its documents: the keys,
# the prose, and these tools. A name-shaped string is as much of a finding in
# ground_truth.json as it is in a PDF.
TEXT_SUFFIXES = {".json", ".csv", ".md", ".py"}

# Capitalised-word sequences that are places, institutions, headings or
# scientific terms rather than people. Any token appearing here disqualifies
# the candidate.
NON_PERSON_TOKENS = {
    # institutions / places invented by the generator
    "Coastal", "State", "University", "Institute", "Technology", "Ridgewater",
    "Calder", "Falls", "Meridian", "Ocean", "Harborview", "Bay", "Marine",
    "Laboratory", "Laboratories", "Center", "Centre", "College", "School",
    "Department", "Division", "Office", "Foundation", "Agency", "Bureau",
    "Administration", "Service", "Survey", "Station", "Facility", "Campus",
    "Building", "Hall", "Room", "Estuary", "River", "Harbor", "Harbour",
    "Island", "Sound", "Coast", "Atlantic", "Pacific", "Gulf", "North",
    "South", "East", "West", "National", "Federal", "Regional", "Council",
    "Board", "Committee", "Program", "Project", "Section", "Appendix",
    "Table", "Figure", "Page", "Year", "Month", "Volume", "Chapter", "Part",
    # sponsors and standards bodies
    "NSF", "NOAA", "EPA", "NIH", "USDA", "USGS", "ISO", "ASTM", "FDA",
    "Science", "Foundation", "Oceanic", "Atmospheric", "Environmental",
    "Protection", "Geological", "Health", "Agriculture",
    # document / proposal vocabulary
    "Budget", "Justification", "Summary", "Description", "Data", "Management",
    "Plan", "References", "Cited", "Facilities", "Equipment", "Resources",
    "Postdoc", "Postdoctoral", "Mentoring", "Biographical", "Sketches",
    "Current", "Pending", "Support", "Rate", "Agreement", "Policy", "Indirect",
    "Direct", "Costs", "Cost", "Total", "Base", "Salary", "Fringe", "Benefits",
    "Travel", "Participant", "Tuition", "Remission", "Subaward", "Award",
    "Modified", "Distribution", "Broader", "Impacts", "Intellectual", "Merit",
    "Results", "Prior", "Methods", "Approach", "Objective", "Objectives",
    "Aim", "Aims", "Task", "Phase", "Quality", "Control", "Assurance",
    "Principal", "Investigator", "Senior", "Personnel", "Graduate", "Student",
    "Research", "Associate", "Assistant", "Professor", "Postdoctoral",
    "Scholar", "Fellow", "Director", "Provost", "President", "Vice",
    "Negotiator", "Officer", "Official", "Representative", "Signature",
    "Date", "Name", "Title", "Role", "Effort", "Months", "Person",
    "Synthetic", "Benchmark", "Fictional", "Valid", "Proposal",
    "Institutional", "Record", "Not", "This", "The", "All",
    # science vocabulary that capitalises mid-sentence
    "Imaging", "Flow", "Cytometer", "Cytometry", "Species", "Taxon", "Taxa",
    "Phytoplankton", "Zooplankton", "Chlorophyll", "Salinity", "Temperature",
    "Sampling", "Field", "Site", "Sites", "Season", "Spring", "Summer",
    "Autumn", "Fall", "Winter", "Model", "Models", "Training", "Validation",
    "Detection", "Limit", "Limits", "Assay", "Sample", "Samples", "Deep",
    "Learning", "Machine", "Neural", "Network", "Networks", "Classifier",
    "Monitoring", "Ecosystem", "Bloom", "Blooms", "Harmful", "Algal",
    "Biological", "Oceanography", "Ecological", "Metadata", "Language",
    "Genomics", "Core", "Renewable", "Energy", "Electrical", "Engineering",
    # section headings that read as First-Last
    "Sponsored", "Programs", "Synergistic", "Activities", "Professional",
    "Preparation", "World", "Register", "Success", "Criteria", "Scientific",
    "Background", "Preliminary", "Studies", "Expected", "Outcomes",
    "Decision", "Gates", "New",
}

# Honorifics strongly imply a person follows.
HONORIFIC = re.compile(r"\b(?:Dr|Prof|Professor|Mr|Ms|Mrs|Mx)\.?\s+")

# First Last, or First M. Last, or First Middle Last.
#
# The separators are `[ \t]+`, not `\s+`, because a name must not be allowed to
# span a line break. `docx_lines()` yields a whole table cell as one string,
# newlines and all, so with `\s+` a signature block reading
# "<name>\nVice President for Research" matched First-Middle-Last *across* the
# break, landed the third group on "Vice", and was then discarded as a
# non-person token — the name went unreported. That is exactly the v0.3.2
# signature-block shape this scanner exists to catch, and the failure got
# quieter as NON_PERSON_TOKENS grew.
NAME = re.compile(
    r"\b([A-Z][a-z]{2,})[ \t]+(?:([A-Z]\.|[A-Z][a-z]{2,})[ \t]+)?([A-Z][a-z]{2,}(?:-[A-Z][a-z]{2,})?)\b"
)


def pdf_pages(path):
    with fitz.open(path) as doc:
        for number, page in enumerate(doc, start=1):
            yield number, page.get_text()


def table_lines(table):
    """Every cell of a table, including the cells of tables nested inside it.

    `_Cell.text` joins the cell's own paragraphs and stops there, so a table
    laid out inside a table cell — the person-months grid in the Current and
    Pending documents is one — contributes nothing unless the walk recurses.
    A name in a nested cell is exactly the shape this scanner exists to catch,
    and it was invisible until this recursion existed.
    """
    for row in table.rows:
        for cell in row.cells:
            yield cell.text
            for nested in cell.tables:
                yield from table_lines(nested)


def docx_lines(path):
    doc = Document(path)
    for paragraph in doc.paragraphs:
        yield paragraph.text
    for table in doc.tables:
        yield from table_lines(table)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph.text
        for paragraph in section.footer.paragraphs:
            yield paragraph.text


def _element_text(element):
    """The concatenated `<t>` runs under an element, at any depth."""
    return "".join(node.text or "" for node in element.iterfind(".//{*}t"))


def _sheet_names(archive):
    """Worksheet part name -> the sheet name a reader sees, best effort."""
    names = {}
    try:
        rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        target_for = {rel.get("Id"): rel.get("Target") or ""
                      for rel in rels.iterfind("{*}Relationship")}
        book = ET.fromstring(archive.read("xl/workbook.xml"))
        rid = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        for sheet in book.iterfind("{*}sheets/{*}sheet"):
            target = target_for.get(sheet.get(rid), "").lstrip("/")
            part = target if target.startswith("xl/") else f"xl/{target}"
            names[part] = sheet.get("name") or part
    except (KeyError, ET.ParseError):
        pass
    return names


def xlsx_cells(path):
    """Every cell's text in a workbook, as (sheet, text).

    Read with `zipfile` rather than openpyxl deliberately: this tool already
    ships alongside the corpus and is run from wherever it was unpacked, so it
    stays on the standard library plus the two readers it cannot do without.
    A workbook is a zip of XML — a cell holds an index into the shared-string
    table, an inline string, or a literal value, and a formula's text can hide
    a name as easily as a label can.

    Until this existed the `.xlsx` branch yielded no lines at all, which made
    the scanner blind to the one file type whose *cells* hold person-shaped
    rows: the COA workbooks.
    """
    with zipfile.ZipFile(path) as archive:
        members = set(archive.namelist())
        shared = []
        if "xl/sharedStrings.xml" in members:
            table = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            shared = [_element_text(item) for item in table.iterfind("{*}si")]
        sheet_names = _sheet_names(archive)
        for part in sorted(name for name in members
                           if name.startswith("xl/worksheets/") and name.endswith(".xml")):
            sheet = sheet_names.get(part, Path(part).stem)
            root = ET.fromstring(archive.read(part))
            for cell in root.iterfind(".//{*}c"):
                kind = cell.get("t")
                if kind == "s":
                    value = cell.findtext("{*}v")
                    index = int(value) if value and value.isdigit() else None
                    text = shared[index] if index is not None and index < len(shared) else ""
                elif kind == "inlineStr":
                    text = _element_text(cell)
                else:
                    text = cell.findtext("{*}v") or ""
                formula = cell.findtext("{*}f") or ""
                for chunk in (text, formula):
                    if chunk.strip():
                        yield sheet, chunk


def office_metadata(path):
    """Author/creator fields, where a personal name hides from text extraction."""
    found = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("docProps/"):
                continue
            blob = archive.read(name).decode("utf-8", errors="ignore")
            for tag in ("dc:creator", "cp:lastModifiedBy", "dc:title",
                        "cp:lastPrinted", "dc:description"):
                for value in re.findall(rf"<{tag}>(.*?)</{tag}>", blob, re.S):
                    if value.strip():
                        found.append(f"{name}:{tag}={value.strip()}")
    return found


def pdf_metadata(path):
    with fitz.open(path) as doc:
        return [f"{key}={value}" for key, value in (doc.metadata or {}).items()
                if value and key in {"author", "title", "subject", "creator", "producer"}]


def candidates(line):
    """Name-shaped substrings in a line, minus the obvious non-people."""
    hits = []
    for match in NAME.finditer(line):
        parts = [part for part in match.groups() if part]
        if any(part.rstrip(".") in NON_PERSON_TOKENS for part in parts):
            continue
        hits.append(match.group(0))
    for match in HONORIFIC.finditer(line):
        tail = line[match.end():match.end() + 40].split(",")[0].strip()
        if tail and tail[0].isupper():
            hits.append(f"{match.group(0).strip()} {tail}")
    return hits


def is_permitted(path: Path, relative: Path) -> bool:
    if path.name in PERMITTED_FILES:
        return True
    posix = relative.as_posix()
    if any(posix == permitted or posix.endswith("/" + permitted)
           for permitted in PERMITTED_TREE_FILES):
        return True
    return any(fnmatch.fnmatch(path.name, pattern) for pattern in PERMITTED_TREE_GLOBS)


def collect(package: Path, show_permitted: bool):
    """Every candidate outside a permitted file, as (file, where, hit, context)."""
    findings = []
    permitted_count = 0
    scanned = 0

    for path in sorted(package.rglob("*")):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix not in {".pdf", ".docx", ".xlsx"} | TEXT_SUFFIXES:
            continue
        scanned += 1
        relative = path.relative_to(package)
        permitted = is_permitted(path, relative)

        if suffix == ".pdf":
            lines = ((f"p.{number}", line)
                     for number, text in pdf_pages(path)
                     for line in text.splitlines())
            meta = pdf_metadata(path)
        elif suffix == ".docx":
            lines = (("body", line) for line in docx_lines(path))
            meta = office_metadata(path)
        elif suffix == ".xlsx":
            lines = ((sheet, text) for sheet, text in xlsx_cells(path))
            meta = office_metadata(path)
        else:
            # Keys, prose and tooling. `where` deliberately does not carry a
            # line number: the baseline would then churn on every unrelated
            # edit above the hit, and the context line is what a reviewer
            # actually reads.
            lines = (("text", line)
                     for line in path.read_text(encoding="utf-8",
                                                errors="ignore").splitlines())
            meta = []

        for value in meta:
            # One line at a time. `office_metadata()` reads with `re.S`, so a
            # wrapped `dc:description` comes back multi-line, and the honorific
            # pass's 40-character tail would run straight across the break into
            # a hit no reviewer can read. No metadata value in the corpus is
            # multi-line today, so this leaves every current finding alone; the
            # recorded context is still the whole value.
            for line in value.splitlines():
                for hit in candidates(line):
                    findings.append((str(relative), "metadata", hit, value[:100]))

        for where, line in lines:
            if any(marker in line for marker in CREDIT_MARKERS):
                continue
            hits = candidates(line)
            if not hits:
                continue
            if permitted:
                permitted_count += len(hits)
                if show_permitted:
                    for hit in hits:
                        print(f"  permitted  {relative} {where}: {hit}")
                continue
            for hit in hits:
                findings.append((str(relative), where, hit, line.strip()[:100]))

    return sorted(set(findings)), permitted_count, scanned


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("package", type=Path)
    parser.add_argument("--show-permitted", action="store_true",
                        help="also list names found in permitted files")
    parser.add_argument("--baseline", type=Path,
                        help="JSON list of reviewed [file, where, hit, context] findings; "
                             "only findings absent from it fail the run")
    parser.add_argument("--emit-baseline", type=Path,
                        help="write the current findings in baseline form and exit 0 "
                             "(review every line before committing it)")
    args = parser.parse_args()

    findings, permitted_count, scanned = collect(args.package, args.show_permitted)

    print(f"scanned {scanned} files")
    print(f"names in permitted files (References Cited, denylist, baselines): "
          f"{permitted_count}")

    if args.emit_baseline:
        args.emit_baseline.write_text(
            json.dumps([list(finding) for finding in findings], indent=2) + "\n"
        )
        print(f"wrote {len(findings)} finding(s) to {args.emit_baseline} — review them")
        return

    if args.baseline is None:
        if not findings:
            print("PERSON-NAME SCAN: PASS — no name-shaped text outside permitted files")
            return
        print(f"PERSON-NAME SCAN: {len(findings)} candidate(s) to review")
        for relative, where, hit, context in findings:
            print(f"- {relative} [{where}] {hit!r}")
            print(f"    {context}")
        sys.exit(1)

    # A finding is identified by (file, where, hit); the context is recorded so
    # the baseline can be reviewed in a diff rather than taken on faith, but it
    # is not part of the identity — rewording the sentence around a reviewed
    # place name should not read as a new name. One name can therefore occur on
    # several lines of the same page under a single identity, which is why the
    # two counts below need not agree.
    baseline = json.loads(args.baseline.read_text())
    reviewed = {tuple(entry[:3]) for entry in baseline}
    current = {finding[:3] for finding in findings}

    new = [finding for finding in findings if finding[:3] not in reviewed]
    stale = sorted(reviewed - current)

    print(f"baseline: {len(baseline)} reviewed entries, {len(reviewed)} distinct")
    print(f"current : {len(findings)} findings, {len(current)} distinct")
    for relative, where, hit in stale:
        print(f"  warning: baseline entry no longer found — {relative} [{where}] {hit!r}")

    if new:
        print(f"PERSON-NAME SCAN: FAIL — {len(new)} finding(s) not in the baseline")
        for relative, where, hit, context in new:
            print(f"- {relative} [{where}] {hit!r}")
            print(f"    {context}")
        sys.exit(1)

    print("PERSON-NAME SCAN: PASS — every finding is a reviewed baseline entry")


if __name__ == "__main__":
    main()
