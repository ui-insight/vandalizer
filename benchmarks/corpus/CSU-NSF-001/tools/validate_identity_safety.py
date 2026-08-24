#!/usr/bin/env python3
"""Release gate for synthetic identities in CSU-NSF-001.

A denylist of the natural-person names v0.3.2 used, plus the role identifiers
that replaced them. It proves the retired names are gone; it cannot catch a
*new* real-person name, which is what `scan_person_names.py` is for.

Takes one or more directories so the same gate covers both halves of the
release: the keys as they sit in the tree, and the unpacked documents.

Run: uv run --with pypdf --with python-docx python \\
       benchmarks/corpus/CSU-NSF-001/tools/validate_identity_safety.py \\
       benchmarks/corpus/CSU-NSF-001 [DIR ...]
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


# The natural-person names earlier corpus versions used for synthetic roles,
# retired in v0.3.3; the denylist has to spell them out in order to test for them.
LEGACY_NAMES = [
    "Elena Maris", "David Okafor", "Priya Nolen", "Samuel Reed", "Andrew Rhyne",
    "Sedgemoor", "Obuya", "Vasquez-Osei", "Bellweather", "Okonkwo-Hale", "Ryland",
    "Draymore", "Nakagawa-Pruitt", "Anstruther", "Corliss", "Tanaka-Boyd", "Pemberly",
]
REQUIRED_IDS = ["CSU-PI-001", "CSU-COI-001", "CSU-VPR-001", "FED-NEG-001"]
FORBIDDEN_FAKE_CITATION_MARKERS = [
    "(fictional)",
    "All products listed are fictional works",
]

# The biographical sketches carry the synthetic product records, ten to a
# person, and each person's records are their own series. v0.4.0 held all
# twenty in one combined document; v0.5.0 splits them per person, as PAPPG
# 24-1 II.D.2.b requires. The count is asserted per file rather than over the
# corpus as a whole because a document keyed on a filename that no longer
# exists is a check that passes by never running — which is what this
# assertion did between the split and this line being written.
SYN_PUB_RECORDS = {
    "10_Biographical_Sketch_PI.docx": ("SYN-PUB-PI", 10),
    "11_Biographical_Sketch_CoPI.docx": ("SYN-PUB-COI", 10),
}

# The document that says the release half is here at all. Where it is scanned,
# the personnel set below must be too.
RELEASE_SENTINEL = "04_Project_Description.docx"

# Files that must be present and scanned wherever the documents are. A release
# that drops one of them silently loses whatever this gate had to say about it,
# and every one of these carries either synthetic person records, person-shaped
# affiliation rows, or both.
REQUIRED_SCAN_FILES = [
    "10_Biographical_Sketch_PI.docx",
    "11_Biographical_Sketch_CoPI.docx",
    "12_Current_Pending_PI.docx",
    "13_Current_Pending_CoPI.docx",
    "14_Synergistic_Activities_PI.docx",
    "15_Synergistic_Activities_CoPI.docx",
    "16_CSU_Research_Infrastructure_Summary.docx",
    "COA_PI.xlsx",
    "COA_CoPI.xlsx",
]


def docx_visible_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def office_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        )


def pdf_text(path: Path) -> str:
    return "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)


def scan(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return docx_visible_text(path) + "\n" + office_xml(path)
    if path.suffix.lower() == ".xlsx":
        return office_xml(path)
    if path.suffix.lower() == ".pdf":
        return pdf_text(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("packages", type=Path, nargs="+",
                        help="one or more directories to scan (keys dir, unpacked assets)")
    args = parser.parse_args()

    failures: list[str] = []
    corpus_text = []
    scanned = 0
    seen_names: set[str] = set()
    product_ids: dict[str, set[str]] = {}

    for package in args.packages:
        files = [
            path for path in package.rglob("*")
            if path.is_file()
            and path.suffix.lower() in {".docx", ".pdf", ".xlsx", ".json", ".csv", ".md"}
        ]
        scanned += len(files)
        for path in sorted(files):
            text = scan(path)
            corpus_text.append(text)
            seen_names.add(path.name)
            for name in LEGACY_NAMES:
                if name in text:
                    failures.append(
                        f"{path.relative_to(package)}: forbidden legacy identity '{name}'"
                    )
            if path.name in SYN_PUB_RECORDS:
                series, expected = SYN_PUB_RECORDS[path.name]
                for marker in FORBIDDEN_FAKE_CITATION_MARKERS:
                    if marker in text:
                        failures.append(
                            f"{path.name}: legacy fictional-citation marker '{marker}'"
                        )
                visible = docx_visible_text(path)
                occurrences = re.findall(r"SYN-PUB-(?:PI|COI)-\d{3}", visible)
                found = {ident for ident in occurrences
                         if ident.startswith(f"{series}-")}
                # Every record id the file carries, whichever series it belongs
                # to: that is what the cross-file check below has to compare,
                # because an id that shows up in the wrong sketch is exactly the
                # id the two sketches end up sharing.
                product_ids[path.name] = set(occurrences)
                if len(found) != expected:
                    failures.append(
                        f"{path.name}: expected {expected} unique {series} records, "
                        f"found {len(found)}"
                    )
                # A count of *unique* ids cannot see a record listed twice. A
                # sketch that grows an eleventh row by copying an existing one
                # still yields ten unique ids, so the count check passes while
                # the products quietly stop being ten distinct products.
                repeated = sorted(ident for ident in set(occurrences)
                                  if occurrences.count(ident) > 1)
                if repeated:
                    failures.append(
                        f"{path.name}: product record id listed more than once: "
                        f"{', '.join(repeated)}"
                    )
                stray = set(occurrences) - found
                if stray:
                    failures.append(
                        f"{path.name}: carries product records from another person's "
                        f"series: {', '.join(sorted(stray))}"
                    )

    # No record id may appear in two sketches: the whole point of splitting the
    # sketches per person is that each person's products are their own. This
    # compares every id each file carries, own series or not — comparing only
    # the ids that matched each file's own series would intersect two
    # prefix-disjoint sets and could never fire. The per-file stray check reads
    # one file at a time and cannot make this statement; it is also the only
    # check left standing if two entries in SYN_PUB_RECORDS are ever keyed to
    # the same series, which would make every shared id look native to both.
    overlapping = set.intersection(*product_ids.values()) if len(product_ids) > 1 else set()
    if overlapping:
        failures.append(
            f"SYN-PUB record ids are shared between biographical sketches: "
            f"{', '.join(sorted(overlapping))}"
        )

    # Only the run that can see the documents can say anything about them, and
    # the presence of the proposal narrative is what says they are there: the
    # key-only tree run scans no Word file at all, so demanding the personnel
    # set there would fail every pull request.
    if RELEASE_SENTINEL in seen_names:
        for required in REQUIRED_SCAN_FILES:
            if required not in seen_names:
                failures.append(f"required file not scanned: {required}")

    all_text = "\n".join(corpus_text)
    for role_id in REQUIRED_IDS:
        if role_id not in all_text:
            failures.append(f"required role identifier missing: {role_id}")
    if failures:
        print("IDENTITY SAFETY: FAIL")
        print("\n".join(f"- {failure}" for failure in failures))
        sys.exit(1)
    print(f"IDENTITY SAFETY: PASS ({scanned} files scanned)")


if __name__ == "__main__":
    main()
